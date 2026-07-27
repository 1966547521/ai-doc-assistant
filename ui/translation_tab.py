"""Translation tab for AI Document Assistant."""
import base64
import io
import streamlit as st
from src.translation_engine import TranslationEngine
from ui.utils import persist_content


def render_translation_tab():
    """Render the document translation tab."""
    st.header("🌍 文档翻译")

    col1, col2 = st.columns(2)

    with col1:
        target_lang = st.selectbox(
            "目标语言",
            options=[lang["code"] for lang in TranslationEngine.SUPPORTED_LANGUAGES],
            format_func=lambda x: next(
                lang["name"] for lang in TranslationEngine.SUPPORTED_LANGUAGES if lang["code"] == x
            ),
            index=1
        )

    with col2:
        translate_mode = st.radio(
            "翻译模式",
            options=["全文翻译", "段落翻译"],
            horizontal=True
        )

    if translate_mode == "全文翻译" and not st.session_state.documents_uploaded:
        st.info("👈 全文翻译需要先上传文档。您也可以切换到「段落翻译」模式，直接粘贴文本翻译。")
        return

    if translate_mode == "段落翻译":
        paragraph_text = st.text_area(
            "输入要翻译的文本",
            height=200,
            placeholder="在此粘贴要翻译的文本...",
            key="translate_input"
        )

    detect_lang = st.checkbox("自动检测源语言（防止误翻）", value=True)

    if st.button("开始翻译"):
        _do_translation(target_lang, translate_mode, paragraph_text if translate_mode == "段落翻译" else None, detect_lang)

    # ── Show saved result with HTML download links ──
    saved = st.session_state.get("_translation_result")
    if saved and not st.session_state.get("_translation_streaming"):
        st.divider()
        st.subheader(f"📄 翻译完成 · {_get_target_name(saved['lang'])}")
        content = saved.get("content", "")
        # Strip progress markers for display
        display_content = _clean_result(content)
        with st.expander("查看翻译", expanded=False):
            st.markdown(display_content)
        _render_html_downloads(display_content, saved["lang"])


def _do_translation(target_lang, translate_mode, paragraph_text, detect_lang):
    if translate_mode == "段落翻译" and not paragraph_text:
        st.warning("请输入要翻译的文本")
        return

    service = st.session_state.application_service
    translator = service.translation_engine
    text_to_translate = paragraph_text if (translate_mode == "段落翻译" and paragraph_text) else service.document_text

    # ── Language detection (cached per text) ──
    if detect_lang and text_to_translate.strip():
        text_key = text_to_translate[:500]
        if st.session_state.get("_detect_cache_key") != text_key:
            with st.spinner("正在检测源语言..."):
                detected = translator.detect_language(text_to_translate)
            st.session_state._detect_cache_key = text_key
            st.session_state._detect_cache_result = detected
        else:
            detected = st.session_state._detect_cache_result

        target_name = _get_target_name(target_lang)
        detected_code = _name_to_code(detected)
        if detected_code == target_lang:
            st.warning(f"⚠️ 源语言「{detected}」与目标语言一致，无需翻译")
            return
        if detected_code:
            st.info(f"🔍 {detected} → {target_name}")
            source_lang = detected_code
        else:
            source_lang = "zh"
    else:
        source_lang = "zh"

    # ── Translate with progress ──
    st.session_state._translation_streaming = True

    is_long = len(text_to_translate) > 12000
    if is_long:
        progress_bar = st.progress(0, text="准备翻译...")
    else:
        progress_bar = None

    with st.spinner("正在翻译..."):
        stream = translator.stream_translate(text_to_translate, target_lang=target_lang, source_lang=source_lang)
        parts = []
        placeholder = st.empty()
        for i, chunk in enumerate(stream):
            parts.append(chunk)
            placeholder.markdown("".join(parts))
            if progress_bar and i % 3 == 0:
                progress_bar.progress(min(i * 0.01 + 0.05, 0.95), "翻译中...")
        result = "".join(parts)
        if progress_bar:
            progress_bar.empty()

    # Clean progress markers from saved result
    clean = _clean_result(result)
    persist_content(st.session_state.get("current_session_id"), "_translation_result",
                    {"lang": target_lang, "content": clean})
    st.session_state._translation_streaming = False
    st.rerun()


def _clean_result(text: str) -> str:
    """Remove progress/status markers from translation output."""
    lines = text.split("\n")
    cleaned = [line for line in lines if not line.startswith(("⏳", "✅"))]
    return "\n".join(cleaned).strip()


# ──────────────────────────────────────────────
#  HTML download links (bypass Streamlit widgets)
# ──────────────────────────────────────────────

def _make_dl_html(data: bytes, filename: str, mime: str, label: str, color: str = "#4a6fa5") -> str:
    """Build an <a> tag that downloads data and shows visual feedback."""
    b64 = base64.b64encode(data).decode()
    return f'''
    <a href="data:{mime};base64,{b64}" download="{filename}"
       style="display:inline-block;width:100%;text-align:center;
              padding:8px 12px;margin-bottom:6px;
              background:{color};color:#fff;text-decoration:none;
              border-radius:8px;font-weight:600;font-size:0.9rem;
              cursor:pointer;box-sizing:border-box;
              transition: opacity 0.15s ease;"
       onclick="this.style.opacity='0.6';setTimeout(()=>this.style.opacity='1',500);">
      {label}
    </a>'''


def _render_html_downloads(result: str, target_lang: str):
    st.markdown("---")
    st.subheader("📥 点击即可下载")

    target_name = _get_target_name(target_lang)
    base_filename = f"translated_{target_lang}"

    # Pre-compute all formats
    links = []

    # TXT (always available)
    txt_data = result.encode("utf-8")
    links.append(_make_dl_html(txt_data, f"{base_filename}.txt", "text/plain",
                                "📄 纯文本 (.txt)", "#4a6fa5"))

    # MD (always available)
    md_content = f"# 翻译结果\n> 目标语言: {target_name}\n\n---\n\n{result}\n"
    links.append(_make_dl_html(md_content.encode("utf-8"), f"{base_filename}.md", "text/markdown",
                                "📝 Markdown (.md)", "#447e5a"))

    # DOCX
    docx = _generate_docx(result, target_name)
    if docx:
        links.append(_make_dl_html(docx, f"{base_filename}.docx",
                                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    "📗 Word (.docx)", "#b8753a"))
    else:
        links.append('<div style="color:#889;font-size:0.85rem;text-align:center;padding:8px;">📗 Word 不可用 (需 python-docx)</div>')

    # PDF
    pdf = _generate_pdf(result, target_name)
    if pdf:
        links.append(_make_dl_html(pdf, f"{base_filename}.pdf", "application/pdf",
                                    "📘 PDF (.pdf)", "#b54444"))
    else:
        links.append('<div style="color:#889;font-size:0.85rem;text-align:center;padding:8px;">📘 PDF 不可用 (需 fpdf2 + 中文字体)</div>')

    # Render in a 2-column grid
    col1, col2 = st.columns(2)
    with col1:
        st.markdown(links[0], unsafe_allow_html=True)
        st.markdown(links[2], unsafe_allow_html=True)
    with col2:
        st.markdown(links[1], unsafe_allow_html=True)
        st.markdown(links[3], unsafe_allow_html=True)


# ── Helpers ──

LANG_CODE_MAP = {"中文": "zh", "English": "en", "日本語": "ja", "한국어": "ko",
                 "Français": "fr", "Deutsch": "de", "Español": "es", "Русский": "ru"}


def _name_to_code(name: str) -> str:
    return LANG_CODE_MAP.get(name, "")


def _get_target_name(code: str) -> str:
    for lang in TranslationEngine.SUPPORTED_LANGUAGES:
        if lang["code"] == code:
            return lang["name"]
    return code


def _generate_docx(text: str, lang_name: str) -> bytes | None:
    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt
    except ImportError:
        return None
    doc = DocxDocument()
    doc.add_heading(f"翻译结果 — {lang_name}", level=1)
    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    style.font.name = "Arial"
    for line in text.split("\n"):
        cleaned = line.strip().replace("**", "").replace("*", "").replace("`", "")
        doc.add_paragraph(cleaned) if cleaned else doc.add_paragraph("")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def _generate_pdf(text: str, lang_name: str) -> bytes | None:
    """Generate PDF using ReportLab (robust CJK support via TTFont)."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib.units import mm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    except ImportError:
        return None

    path = _find_cjk_font()
    if not path:
        return None

    from xml.sax.saxutils import escape

    try:
        font_name = "CJKFont"
        pdfmetrics.registerFont(TTFont(font_name, path))

        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4,
                                leftMargin=15*mm, rightMargin=15*mm,
                                topMargin=10*mm, bottomMargin=15*mm)

        style = ParagraphStyle("Body", fontName=font_name, fontSize=11,
                                leading=16, spaceAfter=2)
        title = ParagraphStyle("Title", fontName=font_name, fontSize=14,
                                leading=20, spaceAfter=8)

        elements = []
        elements.append(Paragraph(f"Translation — {escape(lang_name)}", title))
        elements.append(Spacer(1, 4*mm))

        for para in text.split("\n"):
            line = para.strip()
            if not line:
                elements.append(Spacer(1, 4*mm))
                continue
            # Remove markdown formatting, then XML-escape for ReportLab
            cleaned = line.replace("**", "").replace("*", "").replace("`", "").replace("#", "")
            cleaned = escape(cleaned)
            elements.append(Paragraph(cleaned, style))

        doc.build(elements)
        buf.seek(0)
        return buf.getvalue()
    except Exception:
        return None


def _find_cjk_font() -> str | None:
    import os
    import sys
    if sys.platform == "win32":
        fd = os.path.join(os.environ.get("WINDIR", r"C:\Windows"), "Fonts")
        for f in ["simhei.ttf", "simkai.ttf", "msyh.ttc", "simsun.ttc"]:
            p = os.path.join(fd, f)
            if os.path.exists(p):
                return p
    else:
        for p in ["/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
                   "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
                   "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc"]:
            if os.path.exists(p):
                return p
    return None
